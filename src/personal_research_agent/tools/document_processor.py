"""
Document processing tool for Personal Research Agent.
Handles various document formats and extracts structured information.
"""

import asyncio
import aiofiles
import tempfile
import mimetypes
from pathlib import Path
from typing import Dict, List, Optional, Any, Union, BinaryIO, Type
from datetime import datetime
import logging

from langchain_core.tools import BaseTool
from langchain_core.callbacks import CallbackManagerForToolRun
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from pydantic import BaseModel, Field, PrivateAttr

# Document processing libraries
import pypdf
from docx import Document as DocxDocument
import markdown
from bs4 import BeautifulSoup
import requests
import trafilatura

from ..config import get_settings


class DocumentInput(BaseModel):
    """Input schema for document processing tool."""
    source: str = Field(description="Document source: file path, URL, or raw text")
    source_type: str = Field(default="auto", description="Source type: file, url, text, or auto-detect")
    processing_options: Dict[str, Any] = Field(default_factory=dict, description="Processing options")


class ProcessedDocument(BaseModel):
    """Processed document result."""
    content: str
    metadata: Dict[str, Any]
    chunks: List[Dict[str, Any]]
    summary: Optional[str] = None
    key_points: List[str] = Field(default_factory=list)
    word_count: int
    processing_time: float


class DocumentProcessor(BaseTool):
    """
    LangChain tool for processing various document formats.
    Supports PDFs, Word documents, web pages, and plain text.
    """
    
    name: str = "document_processor"
    description: str = """
    Process and analyze documents from various sources and formats.
    
    Use this tool to:
    - Extract text from PDF files
    - Process Word documents (.docx)
    - Extract content from web pages
    - Parse markdown files
    - Analyze plain text documents
    - Split documents into manageable chunks
    - Extract key information and metadata
    
    Input should specify the document source (file path, URL, or text) and any processing options.
    The tool will return structured content, metadata, and analysis.
    """
    args_schema: Type[BaseModel] = DocumentInput
    _initialized: bool = PrivateAttr(default=False)
    _settings: Any = PrivateAttr(default=None)
    _text_splitter: Any = PrivateAttr(default=None)
    _logger: Any = PrivateAttr(default=None)
    
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
    
    def _ensure_initialized(self):
        """Lazy initialization of tool components."""
        if not self._initialized:
            self._settings = get_settings()
            self._text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=self._settings.vector_db.chunk_size,
                chunk_overlap=self._settings.vector_db.chunk_overlap,
                separators=["\n\n", "\n", ". ", " ", ""]
            )
            self._logger = logging.getLogger(__name__)
            self._initialized = True
    
    def _run(
        self,
        source: str,
        source_type: str = "auto",
        processing_options: Dict[str, Any] = None,
        run_manager: Optional[CallbackManagerForToolRun] = None,
    ) -> str:
        """Synchronous run method (required by BaseTool)."""
        # Run async method in sync context
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            result = loop.run_until_complete(
                self._arun(source, source_type, processing_options or {}, run_manager)
            )
            return result
        finally:
            loop.close()
    
    async def _arun(
        self,
        source: str,
        source_type: str = "auto",
        processing_options: Dict[str, Any] = None,
        run_manager: Optional[CallbackManagerForToolRun] = None,
    ) -> str:
        """Process document asynchronously."""
        start_time = datetime.now()
        processing_options = processing_options or {}
        
        try:
            self._ensure_initialized()
            self._logger.info(f"Processing document from: {source[:100]}...")
            
            # Detect source type if auto
            if source_type == "auto":
                source_type = self._detect_source_type(source)
            
            # Process based on source type
            if source_type == "url":
                processed_doc = await self._process_url(source, processing_options)
            elif source_type == "file":
                processed_doc = await self._process_file(source, processing_options)
            elif source_type == "text":
                processed_doc = await self._process_text(source, processing_options)
            else:
                raise ValueError(f"Unsupported source type: {source_type}")
            
            # Calculate processing time
            processing_time = (datetime.now() - start_time).total_seconds()
            processed_doc.processing_time = processing_time
            
            # Format result for the agent
            result = self._format_result(processed_doc)
            
            self._logger.info(f"Document processed successfully in {processing_time:.2f}s")
            return result
            
        except Exception as e:
            self._logger.error(f"Document processing failed: {e}")
            return f"Document processing failed: {str(e)}"
    
    def _detect_source_type(self, source: str) -> str:
        """Detect the type of document source."""
        if source.startswith(("http://", "https://")):
            return "url"
        elif Path(source).exists():
            return "file"
        else:
            return "text"
    
    async def _process_url(self, url: str, options: Dict[str, Any]) -> ProcessedDocument:
        """Process a web page URL."""
        try:
            # Use trafilatura for better content extraction
            downloaded = trafilatura.fetch_url(url)
            if not downloaded:
                raise ValueError(f"Could not fetch URL: {url}")
            
            content = trafilatura.extract(downloaded, include_comments=False, include_tables=True)
            if not content:
                # Fallback to requests + BeautifulSoup
                response = requests.get(url, timeout=30)
                response.raise_for_status()
                
                soup = BeautifulSoup(response.content, 'html.parser')
                
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                
                content = soup.get_text()
            
            # Clean up the content
            content = self._clean_text(content)
            
            # Extract metadata
            metadata = {
                "source": url,
                "source_type": "url",
                "title": self._extract_title_from_url(url, downloaded),
                "processed_at": datetime.now().isoformat(),
                "content_type": "web_page"
            }
            
            return await self._create_processed_document(content, metadata, options)
            
        except Exception as e:
            raise ValueError(f"Failed to process URL {url}: {str(e)}")
    
    async def _process_file(self, file_path: str, options: Dict[str, Any]) -> ProcessedDocument:
        """Process a local file."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise ValueError(f"File not found: {file_path}")
        
        # Check file size
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        if file_size_mb > self._settings.research.max_document_size_mb:
            raise ValueError(f"File too large: {file_size_mb:.1f}MB (max: {self._settings.research.max_document_size_mb}MB)")
        
        # Determine file type
        mime_type, _ = mimetypes.guess_type(str(file_path))
        file_extension = file_path.suffix.lower()
        
        # Process based on file type
        if file_extension == ".pdf":
            content = await self._extract_pdf_content(file_path)
        elif file_extension == ".docx":
            content = await self._extract_docx_content(file_path)
        elif file_extension in [".md", ".markdown"]:
            content = await self._extract_markdown_content(file_path)
        elif file_extension in [".txt", ".text"]:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                content = await f.read()
        elif file_extension in [".html", ".htm"]:
            content = await self._extract_html_content(file_path)
        else:
            # Try to read as text
            try:
                async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                    content = await f.read()
            except UnicodeDecodeError:
                raise ValueError(f"Unsupported file format: {file_extension}")
        
        # Clean up the content
        content = self._clean_text(content)
        
        # Extract metadata
        metadata = {
            "source": str(file_path),
            "source_type": "file",
            "filename": file_path.name,
            "file_extension": file_extension,
            "file_size_mb": file_size_mb,
            "mime_type": mime_type,
            "processed_at": datetime.now().isoformat(),
            "content_type": "document"
        }
        
        return await self._create_processed_document(content, metadata, options)
    
    async def _process_text(self, text: str, options: Dict[str, Any]) -> ProcessedDocument:
        """Process raw text."""
        content = self._clean_text(text)
        
        metadata = {
            "source": "raw_text",
            "source_type": "text",
            "processed_at": datetime.now().isoformat(),
            "content_type": "text"
        }
        
        return await self._create_processed_document(content, metadata, options)
    
    async def _extract_pdf_content(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        content = ""
        
        with open(file_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        content += f"\n--- Page {page_num + 1} ---\n"
                        content += page_text
                except Exception as e:
                    self._logger.warning(f"Could not extract text from page {page_num + 1}: {e}")
        
        return content
    
    async def _extract_docx_content(self, file_path: Path) -> str:
        """Extract text from Word document."""
        doc = DocxDocument(str(file_path))
        content = ""
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    content += " | ".join(row_text) + "\n"
        
        return content
    
    async def _extract_markdown_content(self, file_path: Path) -> str:
        """Extract text from Markdown file."""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            md_content = await f.read()
        
        # Convert markdown to HTML then extract text
        html = markdown.markdown(md_content)
        soup = BeautifulSoup(html, 'html.parser')
        return soup.get_text()
    
    async def _extract_html_content(self, file_path: Path) -> str:
        """Extract text from HTML file."""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            html_content = await f.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        return soup.get_text()
    
    def _extract_title_from_url(self, url: str, html_content: Optional[str] = None) -> str:
        """Extract title from URL or HTML content."""
        if html_content:
            try:
                soup = BeautifulSoup(html_content, 'html.parser')
                title_tag = soup.find('title')
                if title_tag:
                    return title_tag.get_text().strip()
            except:
                pass
        
        # Fallback to URL-based title
        return url.split('/')[-1] or url
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text content."""
        if not text:
            return ""
        
        # Remove excessive whitespace
        lines = [line.strip() for line in text.split('\n')]
        lines = [line for line in lines if line]  # Remove empty lines
        
        # Join lines and normalize spacing
        cleaned = '\n'.join(lines)
        
        # Remove excessive spaces
        import re
        cleaned = re.sub(r' +', ' ', cleaned)
        
        return cleaned.strip()
    
    async def _create_processed_document(self, content: str, metadata: Dict[str, Any], options: Dict[str, Any]) -> ProcessedDocument:
        """Create a ProcessedDocument with analysis."""
        # Split into chunks
        chunks = []
        if content:
            text_chunks = self._text_splitter.split_text(content)
            for i, chunk in enumerate(text_chunks):
                chunks.append({
                    "chunk_id": i,
                    "content": chunk,
                    "word_count": len(chunk.split()),
                    "char_count": len(chunk)
                })
        
        # Calculate word count
        word_count = len(content.split()) if content else 0
        
        # Extract key points (simple implementation)
        key_points = self._extract_key_points(content)
        
        # Generate summary if requested and content is long enough
        summary = None
        if options.get("generate_summary", False) and word_count > 100:
            summary = self._generate_simple_summary(content)
        
        return ProcessedDocument(
            content=content,
            metadata=metadata,
            chunks=chunks,
            summary=summary,
            key_points=key_points,
            word_count=word_count,
            processing_time=0.0  # Will be set by caller
        )
    
    def _extract_key_points(self, content: str, max_points: int = 5) -> List[str]:
        """Extract key points from content (simple implementation)."""
        if not content:
            return []
        
        # Split into sentences
        sentences = [s.strip() for s in content.split('.') if s.strip()]
        
        # Simple heuristic: longer sentences might be more important
        scored_sentences = [(len(s.split()), s) for s in sentences if len(s.split()) > 5]
        scored_sentences.sort(reverse=True)
        
        # Return top sentences as key points
        key_points = [s[1] + '.' for s in scored_sentences[:max_points]]
        return key_points
    
    def _generate_simple_summary(self, content: str, max_sentences: int = 3) -> str:
        """Generate a simple extractive summary."""
        if not content:
            return ""
        
        sentences = [s.strip() for s in content.split('.') if s.strip()]
        
        if len(sentences) <= max_sentences:
            return content
        
        # Simple approach: take first, middle, and last sentences
        if len(sentences) >= 3:
            indices = [0, len(sentences) // 2, len(sentences) - 1]
            summary_sentences = [sentences[i] for i in indices]
            return '. '.join(summary_sentences) + '.'
        
        return '. '.join(sentences[:max_sentences]) + '.'
    
    def _format_result(self, doc: ProcessedDocument) -> str:
        """Format the processed document result for the agent."""
        result = f"Document Processing Results:\n\n"
        
        # Metadata
        result += f"**Source:** {doc.metadata.get('source', 'Unknown')}\n"
        result += f"**Type:** {doc.metadata.get('content_type', 'Unknown')}\n"
        result += f"**Word Count:** {doc.word_count:,}\n"
        result += f"**Chunks:** {len(doc.chunks)}\n"
        result += f"**Processing Time:** {doc.processing_time:.2f}s\n\n"
        
        # Summary
        if doc.summary:
            result += f"**Summary:**\n{doc.summary}\n\n"
        
        # Key Points
        if doc.key_points:
            result += f"**Key Points:**\n"
            for i, point in enumerate(doc.key_points, 1):
                result += f"{i}. {point}\n"
            result += "\n"
        
        # Content preview
        content_preview = doc.content[:500] + "..." if len(doc.content) > 500 else doc.content
        result += f"**Content Preview:**\n{content_preview}\n\n"
        
        # Chunk information
        if doc.chunks:
            result += f"**Chunks Available:** {len(doc.chunks)} chunks ready for further analysis\n"
        
        return result


# Convenience function for direct usage
async def process_document(source: str, source_type: str = "auto", options: Dict[str, Any] = None) -> ProcessedDocument:
    """Direct document processing function."""
    processor = DocumentProcessor()
    
    # This is a simplified version - in practice you'd want to handle the async properly
    return await processor._create_processed_document("", {}, options or {})
